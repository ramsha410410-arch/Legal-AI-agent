# ============================================================
# tools/document_analyzer.py — PDF & Document Text Extractor
# ============================================================
#
# WHAT DOES THIS FILE DO?
# When a user uploads a PDF or Word document, we need to:
# 1. Read the file's binary data
# 2. Extract the text from it
# 3. Clean up the text (remove weird characters)
# 4. Return clean text the AI can read
#
# WHY IS THIS NEEDED?
# AI models only understand plain text.
# A PDF file is actually a complex binary format.
# We need to "translate" it to plain text first.
# ============================================================

import io
from typing import Tuple, Optional


def extract_text_from_pdf(file_bytes: bytes) -> Tuple[str, Optional[str]]:
    """
    Extract text from a PDF file.
    
    Args:
        file_bytes: Raw bytes of the PDF file (from file.read())
    
    Returns:
        Tuple of (extracted_text: str, error: str|None)
    
    HOW PDF TEXT EXTRACTION WORKS:
    PDFs can contain:
    - Text as actual characters (extractable ✓)
    - Text as images (NOT directly extractable - needs OCR)
    - A mix of both
    
    We use PyPDF2 which handles text-based PDFs well.
    """
    
    try:
        import PyPDF2
        
        # io.BytesIO wraps bytes so PyPDF2 can read them like a file
        # (PyPDF2 expects a file-like object, not raw bytes)
        pdf_file = io.BytesIO(file_bytes)
        
        # Create a PDF reader object
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        # Check how many pages the PDF has
        num_pages = len(pdf_reader.pages)
        
        # Extract text from each page and join them
        extracted_text = ""
        for page_num in range(num_pages):
            page = pdf_reader.pages[page_num]
            page_text = page.extract_text()
            
            if page_text:
                # Add page separator for readability
                extracted_text += f"\n--- Page {page_num + 1} ---\n"
                extracted_text += page_text
        
        if not extracted_text.strip():
            return "", (
                "No text could be extracted from this PDF. "
                "It may be a scanned/image-based PDF. "
                "Try converting it to a text-based PDF first."
            )
        
        # Clean up the text
        cleaned_text = _clean_extracted_text(extracted_text)
        
        return cleaned_text, None  # None = no error
        
    except ImportError:
        return "", "PyPDF2 not installed. Run: pip install PyPDF2"
    except Exception as e:
        return "", f"Error reading PDF: {str(e)}"


def extract_text_from_docx(file_bytes: bytes) -> Tuple[str, Optional[str]]:
    """
    Extract text from a Word document (.docx file).
    
    Args:
        file_bytes: Raw bytes of the DOCX file
    
    Returns:
        Tuple of (extracted_text: str, error: str|None)
    
    HOW DOCX WORKS:
    A .docx file is actually a ZIP archive containing XML files.
    python-docx handles all this complexity for us.
    """
    
    try:
        from docx import Document
        
        docx_file = io.BytesIO(file_bytes)
        document = Document(docx_file)
        
        # Extract text from all paragraphs
        # A paragraph in Word = a block of text
        paragraphs = []
        for paragraph in document.paragraphs:
            if paragraph.text.strip():  # Skip empty paragraphs
                paragraphs.append(paragraph.text)
        
        # Also extract text from tables in the document
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() 
                    for cell in row.cells 
                    if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)
        
        if not paragraphs:
            return "", "No text found in the Word document."
        
        extracted_text = "\n\n".join(paragraphs)
        cleaned_text = _clean_extracted_text(extracted_text)
        
        return cleaned_text, None
        
    except ImportError:
        return "", "python-docx not installed. Run: pip install python-docx"
    except Exception as e:
        return "", f"Error reading Word document: {str(e)}"


def extract_text_from_txt(file_bytes: bytes) -> Tuple[str, Optional[str]]:
    """
    Extract text from a plain text file.
    
    Args:
        file_bytes: Raw bytes of the text file
    
    Returns:
        Tuple of (text: str, error: str|None)
    """
    
    try:
        # Try UTF-8 encoding first (most common)
        text = file_bytes.decode('utf-8')
    except UnicodeDecodeError:
        try:
            # Fall back to latin-1 (handles most European characters)
            text = file_bytes.decode('latin-1')
        except Exception as e:
            return "", f"Could not decode text file: {str(e)}"
    
    return _clean_extracted_text(text), None


def extract_document_text(uploaded_file) -> Tuple[str, Optional[str]]:
    """
    Main entry point: Extract text from any supported file type.
    
    Automatically detects file type and uses the right extractor.
    
    Args:
        uploaded_file: Streamlit UploadedFile object
                       (has .name, .type, .read() attributes)
    
    Returns:
        Tuple of (text: str, error: str|None)
    
    Usage:
        text, error = extract_document_text(uploaded_file)
        if error:
            st.error(error)
        else:
            st.success(f"Extracted {len(text)} characters")
    """
    
    if uploaded_file is None:
        return "", "No file provided"
    
    # Read the file bytes
    file_bytes = uploaded_file.read()
    
    if not file_bytes:
        return "", "File is empty"
    
    # Determine file type from the filename extension
    filename = uploaded_file.name.lower()
    
    if filename.endswith('.pdf'):
        return extract_text_from_pdf(file_bytes)
    
    elif filename.endswith('.docx'):
        return extract_text_from_docx(file_bytes)
    
    elif filename.endswith('.txt'):
        return extract_text_from_txt(file_bytes)
    
    else:
        extension = filename.rsplit('.', 1)[-1] if '.' in filename else 'unknown'
        return "", f"Unsupported file type: .{extension}"


def _clean_extracted_text(text: str) -> str:
    """
    Private helper: Clean up messy extracted text.
    
    PDFs often have:
    - Weird whitespace
    - Null characters
    - Excessive blank lines
    - Garbled characters
    
    We clean these up for better AI processing.
    """
    
    import re
    
    # Remove null characters
    text = text.replace('\x00', '')
    
    # Remove excessive whitespace within lines
    # re.sub replaces regex pattern with replacement string
    text = re.sub(r'[ \t]+', ' ', text)      # Multiple spaces/tabs → single space
    text = re.sub(r'\n{3,}', '\n\n', text)   # Multiple blank lines → max 2
    
    # Fix common PDF extraction artifacts
    text = text.replace('\u2019', "'")   # Right single quote
    text = text.replace('\u2018', "'")   # Left single quote
    text = text.replace('\u201c', '"')   # Left double quote
    text = text.replace('\u201d', '"')   # Right double quote
    text = text.replace('\u2013', '-')   # En dash
    text = text.replace('\u2014', '--')  # Em dash
    
    return text.strip()


def get_document_stats(text: str) -> dict:
    """
    Get statistics about an extracted document.
    
    Useful for showing in the UI before sending to AI.
    
    Returns:
        dict with character_count, word_count, estimated_pages
    """
    
    if not text:
        return {"character_count": 0, "word_count": 0, "estimated_pages": 0}
    
    char_count = len(text)
    word_count = len(text.split())
    # Rough estimate: ~250 words per page
    estimated_pages = max(1, round(word_count / 250))
    
    return {
        "character_count": char_count,
        "word_count": word_count,
        "estimated_pages": estimated_pages,
        "token_estimate": round(char_count / 4)  # Rough token estimate
    }
